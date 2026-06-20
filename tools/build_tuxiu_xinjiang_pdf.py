from __future__ import annotations

import re
from pathlib import Path
from xml.sax.saxutils import escape

from docx import Document
from docx.oxml.ns import qn
from PIL import Image as PILImage
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    BaseDocTemplate,
    Frame,
    Image,
    KeepTogether,
    LongTable,
    NextPageTemplate,
    PageBreak,
    PageTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path("/Users/bruce/Developer/work/AI文旅/星河寻境")
SOURCE_DOCX = ROOT / "规划资料/产品规划/图秀中华公益行动_新疆首站实施方案-汇报版.docx"
LOGO = ROOT / "规划资料/LOGO/图秀中华 4.png"
OUT_DIR = ROOT / "output/pdf"
OUTPUT_PDF = OUT_DIR / "图秀中华公益行动_新疆首站实施方案-精排版.pdf"

PAGE_W, PAGE_H = A4
MARGIN_L = 22 * mm
MARGIN_R = 22 * mm
MARGIN_T = 23 * mm
MARGIN_B = 22 * mm
CONTENT_W = PAGE_W - MARGIN_L - MARGIN_R

NAVY = colors.HexColor("#07365A")
DEEP_NAVY = colors.HexColor("#052943")
TEAL = colors.HexColor("#168A97")
LIGHT_TEAL = colors.HexColor("#E9F6F7")
GOLD = colors.HexColor("#C9A35A")
PALE_GOLD = colors.HexColor("#F8F2E4")
INK = colors.HexColor("#253241")
MUTED = colors.HexColor("#687685")
LINE = colors.HexColor("#D9E5E8")
PAPER = colors.HexColor("#FBFCFD")
WHITE = colors.white

FONT_BODY = "TXZ-Songti"
FONT_HEAD = "TXZ-Heiti"


def register_fonts() -> None:
    candidates = [
        (FONT_HEAD, "/System/Library/Fonts/STHeiti Medium.ttc"),
        (FONT_BODY, "/System/Library/Fonts/Supplemental/Songti.ttc"),
    ]
    for name, path in candidates:
        pdfmetrics.registerFont(TTFont(name, path))


def clean_text(text: str) -> str:
    return re.sub(r"\s+", " ", text or "").strip()


def p(text: str, style: ParagraphStyle) -> Paragraph:
    return Paragraph(escape(clean_text(text)), style)


def para_html(text: str, style: ParagraphStyle) -> Paragraph:
    return Paragraph(text, style)


def split_main_heading(text: str) -> tuple[str, str]:
    match = re.match(r"^([一二三四五六七八九十]+)、(.+)$", text)
    if not match:
        return "", text
    return match.group(1), match.group(2)


def is_main_heading(text: str) -> bool:
    return bool(re.match(r"^[一二三四五六七八九十]+、", text))


def is_sub_heading(text: str) -> bool:
    return bool(re.match(r"^（[一二三四五六七八九十]+）", text))


def is_attachment_heading(text: str) -> bool:
    return text.startswith("附件一") or text.startswith("附件二")


def extract_body(doc: Document) -> list[dict]:
    body_items: list[dict] = []
    paragraph_index = 0
    table_index = 0
    first_heading_seen = False
    body_started = False

    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            para = doc.paragraphs[paragraph_index]
            paragraph_index += 1
            text = clean_text(para.text)
            if not text:
                continue
            if not body_started:
                if text == "一、项目背景与重要意义":
                    if first_heading_seen:
                        body_started = True
                    else:
                        first_heading_seen = True
                if not body_started:
                    continue
            body_items.append({"kind": "paragraph", "text": text})
        elif child.tag == qn("w:tbl"):
            table = doc.tables[table_index]
            table_index += 1
            if not body_started:
                continue
            rows = []
            for row in table.rows:
                rows.append([clean_text(cell.text) for cell in row.cells])
            body_items.append({"kind": "table", "rows": rows})
    return body_items


def extract_front_matter(doc: Document) -> dict:
    texts = [clean_text(paragraph.text) for paragraph in doc.paragraphs if clean_text(paragraph.text)]
    title = "“图秀中华”公益行动·新疆首站"
    doc_type = "实施方案"
    subtitle = "权威版图资源与人工智能数字化公益教育融合示范项目"
    partners = []
    version = ""
    date = ""
    for text in texts:
        if text.startswith("合作单位："):
            partners.append(text.replace("合作单位：", ""))
        elif text in {"中国地图出版集团", "新疆地方协作单位"}:
            partners.append(text)
        elif text.startswith("方案版本："):
            version = text.replace("方案版本：", "")
        elif text.startswith("日期："):
            date = text.replace("日期：", "")
    return {
        "title": title,
        "doc_type": doc_type,
        "subtitle": subtitle,
        "partners": partners,
        "version": version,
        "date": date,
    }


def extract_toc(doc: Document) -> list[str]:
    in_toc = False
    items: list[str] = []
    for paragraph in doc.paragraphs:
        text = clean_text(paragraph.text)
        if text == "目 录":
            in_toc = True
            continue
        if in_toc:
            if text == "一、项目背景与重要意义" and items:
                break
            if text:
                items.append(text)
    return items


class BrandDocTemplate(BaseDocTemplate):
    def afterFlowable(self, flowable):
        name = getattr(flowable, "_bookmark_name", None)
        title = getattr(flowable, "_bookmark_title", None)
        level = getattr(flowable, "_bookmark_level", 0)
        if name and title:
            self.canv.bookmarkPage(name)
            self.canv.addOutlineEntry(title, name, level=level, closed=False)


def draw_cover(canvas, doc):
    canvas.saveState()
    canvas.setFillColor(PAPER)
    canvas.rect(0, 0, PAGE_W, PAGE_H, fill=1, stroke=0)
    canvas.setFillColor(NAVY)
    canvas.rect(0, PAGE_H - 18 * mm, PAGE_W, 18 * mm, fill=1, stroke=0)
    canvas.setFillColor(TEAL)
    canvas.rect(0, PAGE_H - 19.4 * mm, PAGE_W * 0.66, 1.4 * mm, fill=1, stroke=0)
    canvas.setFillColor(GOLD)
    canvas.rect(PAGE_W * 0.66, PAGE_H - 19.4 * mm, PAGE_W * 0.34, 1.4 * mm, fill=1, stroke=0)

    canvas.setStrokeColor(colors.HexColor("#DDE8EB"))
    canvas.setLineWidth(0.35)
    for i in range(11):
        y = 53 * mm + i * 12 * mm
        canvas.line(19 * mm, y, PAGE_W - 19 * mm, y)

    canvas.setFillColor(LIGHT_TEAL)
    canvas.circle(PAGE_W - 27 * mm, PAGE_H - 49 * mm, 18 * mm, fill=1, stroke=0)
    canvas.setFillColor(PALE_GOLD)
    canvas.circle(24 * mm, 52 * mm, 21 * mm, fill=1, stroke=0)

    canvas.setStrokeColor(GOLD)
    canvas.setLineWidth(1.2)
    canvas.line(MARGIN_L, 61 * mm, PAGE_W - MARGIN_R, 61 * mm)
    canvas.restoreState()


def draw_content_page(canvas, doc):
    canvas.saveState()
    canvas.setFillColor(WHITE)
    canvas.rect(0, 0, PAGE_W, PAGE_H, fill=1, stroke=0)
    logo_w = 28 * mm
    with PILImage.open(LOGO) as im:
        ratio = im.height / im.width
    canvas.drawImage(str(LOGO), MARGIN_L, PAGE_H - 15.5 * mm, width=logo_w, height=logo_w * ratio, mask="auto")
    canvas.setFont(FONT_HEAD, 7.8)
    canvas.setFillColor(MUTED)
    canvas.drawRightString(PAGE_W - MARGIN_R, PAGE_H - 10.8 * mm, "图秀中华公益行动 · 新疆首站实施方案")
    canvas.setStrokeColor(LINE)
    canvas.setLineWidth(0.5)
    canvas.line(MARGIN_L, PAGE_H - 18 * mm, PAGE_W - MARGIN_R, PAGE_H - 18 * mm)

    canvas.setStrokeColor(LINE)
    canvas.line(MARGIN_L, 14.5 * mm, PAGE_W - MARGIN_R, 14.5 * mm)
    canvas.setFont(FONT_BODY, 8)
    canvas.setFillColor(MUTED)
    canvas.drawString(MARGIN_L, 9.5 * mm, "权威版图资源与人工智能数字化公益教育融合示范项目")
    canvas.setFillColor(NAVY)
    canvas.setFont(FONT_HEAD, 8.5)
    canvas.drawRightString(PAGE_W - MARGIN_R, 9.5 * mm, f"{doc.page - 1:02d}")
    canvas.restoreState()


def make_styles() -> dict[str, ParagraphStyle]:
    base = {
        "CoverKicker": ParagraphStyle(
            "CoverKicker",
            fontName=FONT_HEAD,
            fontSize=10,
            leading=13,
            textColor=GOLD,
            alignment=TA_CENTER,
            spaceAfter=9,
            wordWrap="CJK",
        ),
        "CoverTitle": ParagraphStyle(
            "CoverTitle",
            fontName=FONT_HEAD,
            fontSize=27,
            leading=35,
            textColor=NAVY,
            alignment=TA_CENTER,
            spaceAfter=6,
            wordWrap="CJK",
        ),
        "CoverDocType": ParagraphStyle(
            "CoverDocType",
            fontName=FONT_HEAD,
            fontSize=23,
            leading=30,
            textColor=TEAL,
            alignment=TA_CENTER,
            spaceAfter=14,
            wordWrap="CJK",
        ),
        "CoverSubtitle": ParagraphStyle(
            "CoverSubtitle",
            fontName=FONT_BODY,
            fontSize=13,
            leading=19,
            textColor=INK,
            alignment=TA_CENTER,
            spaceAfter=20,
            wordWrap="CJK",
        ),
        "Meta": ParagraphStyle(
            "Meta",
            fontName=FONT_BODY,
            fontSize=9.4,
            leading=15,
            textColor=INK,
            alignment=TA_LEFT,
            wordWrap="CJK",
        ),
        "GuideTitle": ParagraphStyle(
            "GuideTitle",
            fontName=FONT_HEAD,
            fontSize=20,
            leading=26,
            textColor=NAVY,
            spaceAfter=8,
            wordWrap="CJK",
        ),
        "GuideLead": ParagraphStyle(
            "GuideLead",
            fontName=FONT_BODY,
            fontSize=10.5,
            leading=17,
            textColor=INK,
            spaceAfter=12,
            wordWrap="CJK",
        ),
        "CardTitle": ParagraphStyle(
            "CardTitle",
            fontName=FONT_HEAD,
            fontSize=10,
            leading=14,
            textColor=NAVY,
            alignment=TA_CENTER,
            spaceAfter=4,
            wordWrap="CJK",
        ),
        "CardBody": ParagraphStyle(
            "CardBody",
            fontName=FONT_BODY,
            fontSize=8.2,
            leading=12,
            textColor=INK,
            alignment=TA_CENTER,
            wordWrap="CJK",
        ),
        "TOCNum": ParagraphStyle(
            "TOCNum",
            fontName=FONT_HEAD,
            fontSize=8.8,
            leading=13,
            textColor=TEAL,
            alignment=TA_CENTER,
            wordWrap="CJK",
        ),
        "TOCItem": ParagraphStyle(
            "TOCItem",
            fontName=FONT_BODY,
            fontSize=9,
            leading=13,
            textColor=INK,
            wordWrap="CJK",
        ),
        "SectionNum": ParagraphStyle(
            "SectionNum",
            fontName=FONT_HEAD,
            fontSize=12,
            leading=17,
            textColor=WHITE,
            alignment=TA_CENTER,
            wordWrap="CJK",
        ),
        "SectionTitle": ParagraphStyle(
            "SectionTitle",
            fontName=FONT_HEAD,
            fontSize=15.5,
            leading=22,
            textColor=NAVY,
            wordWrap="CJK",
        ),
        "AttachmentTitle": ParagraphStyle(
            "AttachmentTitle",
            fontName=FONT_HEAD,
            fontSize=15,
            leading=22,
            textColor=NAVY,
            wordWrap="CJK",
        ),
        "SubHeading": ParagraphStyle(
            "SubHeading",
            fontName=FONT_HEAD,
            fontSize=11.5,
            leading=17,
            textColor=TEAL,
            spaceBefore=6,
            spaceAfter=4,
            wordWrap="CJK",
        ),
        "Body": ParagraphStyle(
            "Body",
            fontName=FONT_BODY,
            fontSize=9.9,
            leading=16.3,
            firstLineIndent=18,
            textColor=INK,
            spaceAfter=5.8,
            wordWrap="CJK",
        ),
        "Bullet": ParagraphStyle(
            "Bullet",
            fontName=FONT_BODY,
            fontSize=9.8,
            leading=15.6,
            leftIndent=17,
            bulletIndent=5,
            textColor=INK,
            spaceAfter=4.2,
            wordWrap="CJK",
        ),
        "Caption": ParagraphStyle(
            "Caption",
            fontName=FONT_HEAD,
            fontSize=9.5,
            leading=13,
            textColor=NAVY,
            spaceBefore=8,
            spaceAfter=4,
            wordWrap="CJK",
        ),
        "TableHead": ParagraphStyle(
            "TableHead",
            fontName=FONT_HEAD,
            fontSize=8.1,
            leading=11.2,
            textColor=WHITE,
            alignment=TA_CENTER,
            wordWrap="CJK",
        ),
        "TableCell": ParagraphStyle(
            "TableCell",
            fontName=FONT_BODY,
            fontSize=7.7,
            leading=11.4,
            textColor=INK,
            wordWrap="CJK",
        ),
        "TableCellCenter": ParagraphStyle(
            "TableCellCenter",
            fontName=FONT_HEAD,
            fontSize=7.8,
            leading=11.4,
            textColor=NAVY,
            alignment=TA_CENTER,
            wordWrap="CJK",
        ),
    }
    return base


def cover_story(front: dict, styles: dict[str, ParagraphStyle]) -> list:
    with PILImage.open(LOGO) as im:
        ratio = im.height / im.width
    logo_w = 88 * mm
    partner_text = " ｜ ".join(front["partners"])
    meta_rows = [
        [p("合作单位", styles["Meta"]), p(partner_text, styles["Meta"])],
        [p("方案版本", styles["Meta"]), p(front["version"], styles["Meta"])],
        [p("日期", styles["Meta"]), p(front["date"], styles["Meta"])],
    ]
    meta = Table(meta_rows, colWidths=[22 * mm, CONTENT_W - 22 * mm])
    meta.setStyle(
        TableStyle(
            [
                ("BOX", (0, 0), (-1, -1), 0.7, LINE),
                ("INNERGRID", (0, 0), (-1, -1), 0.35, LINE),
                ("BACKGROUND", (0, 0), (0, -1), LIGHT_TEAL),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 9),
                ("RIGHTPADDING", (0, 0), (-1, -1), 9),
                ("TOPPADDING", (0, 0), (-1, -1), 7),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
            ]
        )
    )
    return [
        Spacer(1, 38 * mm),
        Image(str(LOGO), width=logo_w, height=logo_w * ratio),
        Spacer(1, 15 * mm),
        p("公益行动 · 新疆首站", styles["CoverKicker"]),
        p(front["title"], styles["CoverTitle"]),
        p(front["doc_type"], styles["CoverDocType"]),
        p(front["subtitle"], styles["CoverSubtitle"]),
        Spacer(1, 9 * mm),
        meta,
        Spacer(1, 20 * mm),
        p("权威资源 · 数字赋能 · 边疆首发 · 全国复制", styles["CoverKicker"]),
        NextPageTemplate("content"),
        PageBreak(),
    ]


def guide_story(toc: list[str], styles: dict[str, ParagraphStyle]) -> list:
    story: list = []
    story.append(p("阅读导览", styles["GuideTitle"]))
    story.append(
        p(
            "本版本在保留原实施方案内容的基础上，按对外汇报阅读习惯重建层级、表格和版面，突出项目定位、资源体系、数字化支撑与新疆首站示范路径。",
            styles["GuideLead"],
        )
    )

    cards = [
        ("权威资源", "以出版集团标准地图、图书、地球仪和审核知识库作为内容底座。"),
        ("数字赋能", "通过扫码学习、语音讲解、互动问答、内容更新和后台管理提升使用效率。"),
        ("新疆首站", "在边疆地区形成可展示、可汇报、可传播的公益教育样板。"),
        ("全国复制", "沉淀产品包、项目手册、评估指标和长期公益品牌。"),
    ]
    card_cells = []
    for title, desc in cards:
        card_cells.append([p(title, styles["CardTitle"]), p(desc, styles["CardBody"])])
    cards_table = Table([card_cells], colWidths=[CONTENT_W / 4] * 4)
    cards_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (0, 0), LIGHT_TEAL),
                ("BACKGROUND", (1, 0), (1, 0), PALE_GOLD),
                ("BACKGROUND", (2, 0), (2, 0), LIGHT_TEAL),
                ("BACKGROUND", (3, 0), (3, 0), PALE_GOLD),
                ("BOX", (0, 0), (-1, -1), 0.5, LINE),
                ("INNERGRID", (0, 0), (-1, -1), 0.5, WHITE),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ("TOPPADDING", (0, 0), (-1, -1), 10),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 10),
            ]
        )
    )
    story.extend([cards_table, Spacer(1, 9 * mm), p("目录", styles["GuideTitle"])])

    rows = []
    for idx, item in enumerate(toc, start=1):
        if item.startswith("附件"):
            num = "附"
        else:
            num = f"{idx:02d}"
        rows.append([p(num, styles["TOCNum"]), p(item, styles["TOCItem"])])
    toc_table = Table(rows, colWidths=[18 * mm, CONTENT_W - 18 * mm])
    toc_table.setStyle(
        TableStyle(
            [
                ("BOX", (0, 0), (-1, -1), 0.5, LINE),
                ("INNERGRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#EDF2F4")),
                ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#F1FAFB")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 7),
                ("RIGHTPADDING", (0, 0), (-1, -1), 7),
                ("TOPPADDING", (0, 0), (-1, -1), 4.7),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4.7),
            ]
        )
    )
    story.append(toc_table)
    story.append(PageBreak())
    return story


def section_heading(text: str, styles: dict[str, ParagraphStyle], bookmark_index: int) -> KeepTogether:
    num, title = split_main_heading(text)
    table = Table(
        [[p(num, styles["SectionNum"]), p(title, styles["SectionTitle"])]],
        colWidths=[18 * mm, CONTENT_W - 18 * mm],
    )
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (0, 0), NAVY),
                ("BACKGROUND", (1, 0), (1, 0), colors.HexColor("#F1FAFB")),
                ("BOX", (0, 0), (-1, -1), 0.6, LINE),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ("TOPPADDING", (0, 0), (-1, -1), 8),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
            ]
        )
    )
    table._bookmark_name = f"sec_{bookmark_index}"
    table._bookmark_title = text
    table._bookmark_level = 0
    return KeepTogether([Spacer(1, 5 * mm), table, Spacer(1, 3 * mm)])


def attachment_heading(text: str, styles: dict[str, ParagraphStyle], bookmark_index: int) -> list:
    title = p(text, styles["AttachmentTitle"])
    title._bookmark_name = f"att_{bookmark_index}"
    title._bookmark_title = text
    title._bookmark_level = 0
    rule = Table([[""]], colWidths=[CONTENT_W], rowHeights=[1.4 * mm])
    rule.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, -1), GOLD)]))
    return [PageBreak(), title, rule, Spacer(1, 4 * mm)]


def table_caption(rows: list[list[str]]) -> str:
    if not rows:
        return "表格"
    header = tuple(rows[0])
    captions = {
        ("类别", "建设内容", "主要价值"): "五位一体公益资源体系",
        ("类别", "捐赠内容"): "公益捐赠内容清单",
        ("产品包", "适用场景", "主要内容", "主要价值"): "分层公益产品包",
        ("阶段", "重点任务", "阶段成果"): "实施路径三阶段",
        ("类别", "指标示例", "用途"): "公益评估指标体系",
        ("层级", "主要内容"): "数字化平台总体架构",
        ("交付包", "适用对象", "主要内容"): "标准化交付包",
    }
    return captions.get(header, "重点内容表")


def col_widths(col_count: int) -> list[float]:
    if col_count == 2:
        return [38 * mm, CONTENT_W - 38 * mm]
    if col_count == 3:
        return [29 * mm, 79 * mm, CONTENT_W - 108 * mm]
    if col_count == 4:
        return [22 * mm, 31 * mm, 75 * mm, CONTENT_W - 128 * mm]
    return [CONTENT_W / col_count] * col_count


def make_table(rows: list[list[str]], styles: dict[str, ParagraphStyle]) -> list:
    if not rows:
        return []
    rendered = []
    for row_index, row in enumerate(rows):
        rendered_row = []
        for col_index, cell in enumerate(row):
            style = styles["TableHead"] if row_index == 0 else styles["TableCell"]
            if row_index > 0 and col_index == 0:
                style = styles["TableCellCenter"]
            rendered_row.append(p(cell, style))
        rendered.append(rendered_row)
    table = LongTable(rendered, colWidths=col_widths(len(rows[0])), repeatRows=1, splitByRow=1)
    style_cmds = [
        ("BACKGROUND", (0, 0), (-1, 0), NAVY),
        ("TEXTCOLOR", (0, 0), (-1, 0), WHITE),
        ("BOX", (0, 0), (-1, -1), 0.65, LINE),
        ("INNERGRID", (0, 0), (-1, -1), 0.35, LINE),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
    ]
    for row_index in range(1, len(rows)):
        if row_index % 2 == 1:
            style_cmds.append(("BACKGROUND", (0, row_index), (-1, row_index), colors.HexColor("#FAFCFD")))
        else:
            style_cmds.append(("BACKGROUND", (0, row_index), (-1, row_index), colors.HexColor("#F4FAFB")))
    table.setStyle(TableStyle(style_cmds))
    return [
        KeepTogether([p(table_caption(rows), styles["Caption"]), table]),
        Spacer(1, 5 * mm),
    ]


def body_story(items: list[dict], styles: dict[str, ParagraphStyle]) -> list:
    story: list = []
    section_index = 0
    attachment_index = 0
    bullet_context = ""
    for item in items:
        if item["kind"] == "table":
            story.extend(make_table(item["rows"], styles))
            continue
        text = item["text"]
        if is_attachment_heading(text):
            attachment_index += 1
            story.extend(attachment_heading(text, styles, attachment_index))
            bullet_context = ""
        elif is_main_heading(text):
            section_index += 1
            story.append(section_heading(text, styles, section_index))
            bullet_context = ""
        elif is_sub_heading(text):
            story.append(p(text, styles["SubHeading"]))
            bullet_context = text
        else:
            list_like = (
                any(key in bullet_context for key in ["目标", "价值"])
                and len(text) <= 58
            ) or text.startswith(("对学生而言", "对学校而言", "对公益方而言", "对中图而言", "对地方而言"))
            if list_like:
                story.append(Paragraph(escape(text), styles["Bullet"], bulletText="-"))
            else:
                story.append(p(text, styles["Body"]))
    return story


def build_pdf() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    register_fonts()
    source = Document(SOURCE_DOCX)
    styles = make_styles()
    front = extract_front_matter(source)
    toc = extract_toc(source)
    body = extract_body(source)

    doc = BrandDocTemplate(
        str(OUTPUT_PDF),
        pagesize=A4,
        leftMargin=MARGIN_L,
        rightMargin=MARGIN_R,
        topMargin=MARGIN_T,
        bottomMargin=MARGIN_B,
        title="图秀中华公益行动·新疆首站实施方案",
        author="图秀中华",
        subject="权威版图资源与人工智能数字化公益教育融合示范项目",
    )
    cover_frame = Frame(MARGIN_L, MARGIN_B, CONTENT_W, PAGE_H - MARGIN_T - MARGIN_B, id="cover")
    content_frame = Frame(MARGIN_L, MARGIN_B, CONTENT_W, PAGE_H - MARGIN_T - MARGIN_B, id="content")
    doc.addPageTemplates(
        [
            PageTemplate(id="cover", frames=[cover_frame], onPage=draw_cover),
            PageTemplate(id="content", frames=[content_frame], onPage=draw_content_page),
        ]
    )
    story: list = []
    story.extend(cover_story(front, styles))
    story.extend(guide_story(toc, styles))
    story.extend(body_story(body, styles))
    doc.build(story)
    print(OUTPUT_PDF)


if __name__ == "__main__":
    build_pdf()
